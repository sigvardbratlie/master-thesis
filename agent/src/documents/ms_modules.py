from io import BytesIO
import logging
from langchain_core.documents import Document
from models import VectorStoreMetadata,FileType,  WriteDocx
from docx import Document as DocxDocument
from pptx import Presentation
from .base_module import BaseHandler
from markitdown import MarkItDown


logger = logging.getLogger(__name__)


class DocxHandler(BaseHandler):
    def __init__(self,chunk_size : int = 1000, chunk_overlap : int = 200):
        '''Handler for parsing DOCX documents.
        Args:
            chunk_size (int): The maximum size of each text chunk extracted from the DOCX. (default: Splits by paragraph)
            chunk_overlap (int): The number of characters to overlap between chunks. (default: 200)
        '''
        super().__init__(chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    def parse_docx(self, content: bytes, metadata: dict, force_metadata_model: bool = True) -> tuple[str, dict]:
        try:
            props = DocxDocument(BytesIO(content)).core_properties
            metadata_full = {
                            **metadata,
                            "title": props.title,
                            "creator": props.author,
                            "created_at": props.created.isoformat() if props.created else None,
                            "updated_at": props.modified.isoformat() if props.modified else None,
                            "comments": props.comments,
                            "language": props.language,
                            "file_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        }
        except Exception as e:
            logger.exception(f"❌ Failed to load metadata with python-docx ({metadata.get('filename', 'unknown')})")
            metadata_full = metadata
        final_metadata = VectorStoreMetadata.model_validate(metadata_full).model_dump(mode="json") if force_metadata_model else metadata_full
        full_text = MarkItDown().convert(source=BytesIO(content)).markdown
        return full_text, final_metadata

    def mk_docx(self, docx_data : WriteDocx) -> bytes:
        doc = DocxDocument()
        if docx_data.heading:
            doc.add_heading(docx_data.heading, level=1)
        for para in docx_data.paragraphs:
            doc.add_paragraph(para)
        with BytesIO() as output:
            doc.save(output)
            return output.getvalue()
        
class XlsxHandler(BaseHandler):
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        super().__init__(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    
    def parse_xlsx(self, content: bytes, metadata: dict, force_metadata_model: bool = True) -> tuple[str, dict]:
        # try:
        #     props = (BytesIO(content)).core_properties
        #     metadata_full = {
        #                     **metadata,
        #                     "title": props.title,
        #                     "creator": props.author,
        #                     "created_at": props.created.isoformat() if props.created else None,
        #                     "updated_at": props.modified.isoformat() if props.modified else None,
        #                     "comments": props.comments,
        #                     "language": props.language,
        #                     "file_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        #                 }
        # except Exception as e:
        #     logger.exception(f"❌ Failed to load metadata with python-docx ({metadata.get('filename', 'unknown')})")
        metadata_full = metadata
        final_metadata = VectorStoreMetadata.model_validate(metadata_full).model_dump(mode="json") if force_metadata_model else metadata_full
        full_text = MarkItDown().convert(source=BytesIO(content)).markdown
        return full_text, final_metadata
    
class PptxHandler(BaseHandler):
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        super().__init__(chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    def parse_pptx(self, content: bytes, metadata: dict, force_metadata_model: bool = True) -> tuple[str, dict]:
        try:
            props = Presentation(BytesIO(content)).core_properties
            metadata_full = {
                            **metadata,
                            "title": props.title,
                            "creator": props.author,
                            "created_at": props.created.isoformat() if props.created else None,
                            "updated_at": props.modified.isoformat() if props.modified else None,
                            "comments": props.comments,
                            "language": props.language,
                            "file_type": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
                        }
        except Exception as e:
            logger.exception(f"❌ Failed to load PPTX ({metadata.get('filename', 'unknown')})")
            metadata_full = metadata
       
        final_metadata = VectorStoreMetadata.model_validate(metadata_full).model_dump(mode="json") if force_metadata_model else metadata_full

        full_text = MarkItDown().convert(source=BytesIO(content)).markdown
        return full_text, final_metadata
    
    

